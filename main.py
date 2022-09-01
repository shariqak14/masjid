from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from datetime import datetime, timedelta, date
from calendar import Calendar, day_name, monthrange

from tkinter import *
from tkinter.ttk import *

from api import get_prayer_time, get_lunar_date, get_islamic_year
from constants import LUNAR_MONTHS, GREG_MONTHS

import os


def number_of_days_in_month(year=2021, month=2):
    return monthrange(year, month)[1]


def format_cell(
    cell, center=True, bold=False, italic=False, font_size=11, font="Arial"
):
    run = cell.paragraphs[0].runs[0]

    run.font.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font

    run.italic = italic

    if center:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    month_input, year_input = e1.get(), e2.get()

    ####################################################################################
    # Edit Document
    ####################################################################################

    NUM_DAYS = number_of_days_in_month(int(year_input), GREG_MONTHS[month_input])

    doc = Document("templates/" + str(NUM_DAYS) + "_days.docx")

    table = doc.tables[0]

    month_heading = table.cell(0, 10)
    month_heading.text = month_input.upper() + " " + year_input
    format_cell(month_heading, bold=True, font_size=15, font="Times New Roman")

    month_row_title = table.cell(1, 2)
    month_row_title.text = month_input + " " + year_input + " CE"
    format_cell(month_row_title, bold=True, font_size=11, font="Times New Roman")

    ####################################################################################
    # Insert Date
    ####################################################################################

    uniq_lunar_months = []

    for j in range(3, NUM_DAYS + 3):
        my_date = date(int(year_input), GREG_MONTHS[month_input], j - 2)
        name_of_day = str(day_name[my_date.weekday()])[:3]

        day_col = table.cell(j, 2)
        day_col.text = name_of_day
        format_cell(day_col, bold=True)

        day_no_col = table.cell(j, 3)
        day_no_col.text = str(j - 2)
        format_cell(day_no_col, bold=True)

        lunar_month, lunary_day = get_lunar_date(
            int(year_input), GREG_MONTHS[month_input], j - 2
        )

        lunary_day_col = table.cell(j, 4)

        if j == 3:
            lunary_day_col.text = LUNAR_MONTHS[lunar_month - 1] + " " + str(lunary_day)
            format_cell(lunary_day_col, bold=True, italic=True, font_size=10.5)
        elif lunary_day == 1:
            lunary_day_col.text = (
                LUNAR_MONTHS[lunar_month - 1] + " " + str(lunary_day) + "*"
            )
            format_cell(lunary_day_col, bold=True, italic=True, font_size=10.5)
        else:
            lunary_day_col.text = str(lunary_day)
            format_cell(lunary_day_col)

        if lunar_month not in uniq_lunar_months:
            uniq_lunar_months.append(lunar_month)

    month_1, month_2 = uniq_lunar_months[0], uniq_lunar_months[1]
    lunar_col_title = table.cell(2, 4)
    lunar_col_title.text = (
        LUNAR_MONTHS[month_1 - 1] + "\n" + LUNAR_MONTHS[month_2 - 1] + "\n" + get_islamic_year(month_input, year_input) + " AH"
    )
    format_cell(lunar_col_title, font_size=10.5, bold=True, italic=True)

    ####################################################################################
    # Insert Prayer Times
    ####################################################################################

    prayer_times = get_prayer_time(month=month_input, year=year_input)

    num_of_rows = len(prayer_times)

    prayers = {"Fajr": 5, "Duha": 8, "Dhuhr": 9, "Asr": 11, "Maghrib": 13, "Isha": 15}

    for prayer, col_no in prayers.items():
        times = list(map(lambda x: x[prayer], prayer_times))

        for j in range(3, NUM_DAYS + 3):
            pr_time = table.cell(j, col_no)
            pr_time.text = times[j - 3]
            format_cell(pr_time)

            if prayer == "Maghrib":
                iqamah_time = datetime.strptime(times[j - 3], "%H:%M") + timedelta(
                    minutes=10
                )
                iqamah_time = iqamah_time.strftime("%H:%M")
                iqamah_time = iqamah_time if iqamah_time[0] != "0" else iqamah_time[1:]

                iqamah_cell = table.cell(j, col_no + 1)
                iqamah_cell.text = iqamah_time
                format_cell(iqamah_cell, bold=True)

    if not os.path.exists("calendars"):
        os.makedirs("calendars")

    doc.save("calendars/" + month_input + "_" + year_input + ".docx")


if __name__ == "__main__":
    window = Tk()
    window.title("Masjid Calendar")
    window.geometry("400x275")
    window.eval("tk::PlaceWindow . center")

    lbl1 = Label(window, text="Month")

    current_month = int(datetime.now().strftime("%m"))
    current_year = int(datetime.today().year)

    e1 = Combobox(window, width=15)
    e1["values"] = list(GREG_MONTHS.keys())
    e1.current(current_month - 1)
    e1.grid(row=0, column=1)

    lbl1.place(x=80, y=50)
    e1.place(x=180, y=50)

    lbl2 = Label(window, text="Year")

    dflt = IntVar()
    dflt.set(current_year)

    e2 = Spinbox(window, from_=2020, to=2121, width=15, textvariable=dflt)

    lbl2.place(x=80, y=100)
    e2.place(x=180, y=100)

    b1 = Button(window, text="Submit", command=main, width=34)
    b1.place(x=80, y=150)

    window.mainloop()
